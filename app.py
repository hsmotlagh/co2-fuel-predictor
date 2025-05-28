# === Password-protected Streamlit App ===
import streamlit as st
import pandas as pd
import numpy as np
import joblib
import matplotlib.pyplot as plt
from docx import Document
import sys, os

# --- Auth ---
from streamlit.runtime.scriptrunner import get_script_run_ctx
import toml

def check_password():
    secrets_path = os.path.join(".streamlit", "secrets.toml")
    if not os.path.exists(secrets_path):
        st.error("secrets.toml not found")
        st.stop()

    secrets = toml.load(secrets_path).get("auth", {})
    username = secrets.get("username")
    password = secrets.get("password")

    with st.sidebar:
        st.subheader("🔐 Login Required")
        u = st.text_input("Username")
        p = st.text_input("Password", type="password")
        if st.button("Login"):
            if u == username and p == password:
                st.session_state["authenticated"] = True
            else:
                st.error("Invalid credentials")

    if not st.session_state.get("authenticated"):
        st.stop()

check_password()

# --- Header and UI Layout ---
col1, col2 = st.columns([6, 1])
with col1:
    st.markdown("<h1 style='margin-bottom: 0;'>\ud83d\udea2 XGB CO\u2082 & Fuel-Consumption Predictor</h1>", unsafe_allow_html=True)
with col2:
    st.image("AppIcon.png", width=150)

col3, col4 = st.columns([6, 2])
with col3:
    st.markdown("<div style='text-align: right; font-size: 0.85em;'>Created by Hamid<br>\u00a9 2025 All rights reserved</div>", unsafe_allow_html=True)
with col4:
    with st.expander("\u2139\ufe0f", expanded=False):
        st.markdown("""
        <div style='font-size: 0.75em; line-height: 1.2;'>
        <b>About this App</b><br>
        Project: CO\u2082 & Fuel Prediction<br>
        By: Hamid (PhD)<br>
        Year: 2025<br>
        Maritime Research
        </div>
        """, unsafe_allow_html=True)

# --- Paths ---
if getattr(sys, "frozen", False):
    base_path = sys._MEIPASS
else:
    base_path = os.path.dirname(__file__)

model_path  = os.path.join(base_path, "XGB_ML_model") + os.sep
data_path   = os.path.join(base_path, "ML_Apply copy.xlsx")
output_path = os.path.join(base_path, "ML_Prediction_Results_1.xlsx")
report_path = os.path.join(base_path, "Prediction_Report.docx")

xgb_fc  = joblib.load(f"{model_path}xgb_model_fc_combined.pkl")
xgb_co2 = joblib.load(f"{model_path}xgb_model_co2_combined.pkl")
scaler  = joblib.load(f"{model_path}scaler_combined.pkl")

# --- Ranges and Features ---
RANGES = {
    'Pe (kW)': (300.0, 2200.0),
    'Ship Speed (knots)': (9.0, 15.0),
    'Hull Efficiency (\u03b7H)': (1.0, 1.1),
    'Propeller Open Water Efficiency (\u03b7O)': (0.5, 0.65)
}

FEATURE_COLS = [
    'Pe (kW)', 'n (rpm)', 'T (kN)',
    'Hull Efficiency (\u03b7H)', 'Propeller Open Water Efficiency (\u03b7O)',
    'Delivered Power (Pd) (kW)', 'Brake Power (PB) (kW)', 'Ship Speed (knots)',
    'Physics-Based FC','Physics-Based CO2',
    'Ship Speed^2','Ship Speed^3',
    'Pd * Hull Efficiency','Propeller Efficiency * n'
]

sheet_names = [
    "Original Scenar_Coeff", "Paint (5%) Scen_Coeff", "Advance Propell_Coeff",
    "Fin (2%-4%) Sce_Coeff", "Bulbous Bow Sce_Coeff"
]
excel_data = pd.read_excel(data_path, sheet_name=sheet_names)
combined_data = pd.concat(excel_data.values(), ignore_index=True)

# --- Predict Function ---
def predict_from_inputs(vals):
    df = combined_data
    dist = np.sqrt(
        ((df['Pe (kW)']-vals['Pe (kW)'])/RANGES['Pe (kW)'][1])**2 +
        ((df['Ship Speed (knots)']-vals['Ship Speed (knots)'])/RANGES['Ship Speed (knots)'][1])**2
    )
    idx = dist.nsmallest(3).index
    w = 1/(dist.loc[idx]+1e-6); w/=w.sum()
    nrpm = (df.loc[idx,'n (rpm)']*w).sum()
    tkn  = (df.loc[idx,'T (kN)']*w).sum()
    pd_kw, pb_kw = vals['Pe (kW)'], vals['Pe (kW)']*0.95

    row = {
        **vals,
        'n (rpm)': nrpm,
        'T (kN)':  tkn,
        'Delivered Power (Pd) (kW)': pd_kw,
        'Brake Power (PB) (kW)': pb_kw,
        'Physics-Based FC': pd_kw*vals['Hull Efficiency (\u03b7H)']*0.8,
        'Physics-Based CO2': pd_kw*vals['Propeller Open Water Efficiency (\u03b7O)']*0.5,
        'Ship Speed^2': vals['Ship Speed (knots)']**2,
        'Ship Speed^3': vals['Ship Speed (knots)']**3,
        'Pd * Hull Efficiency': pd_kw*vals['Hull Efficiency (\u03b7H)'],
        'Propeller Efficiency * n': nrpm*vals['Propeller Open Water Efficiency (\u03b7O)']
    }
    X = pd.DataFrame([row])[FEATURE_COLS]
    Xs = scaler.transform(X)
    return float(xgb_fc.predict(Xs)[0]), float(xgb_co2.predict(Xs)[0])

# --- Plot Function ---
def make_plots(user_speed, fc_pred, co2_pred):
    fig, (ax1,ax2) = plt.subplots(2,1, figsize=(6,6))
    speeds = sorted(combined_data['Ship Speed (knots)'].unique())
    for name, df in excel_data.items():
        ax1.plot(df['Ship Speed (knots)'], df['Final Predicted Fuel Consumption XGB After MC (kg/h)'], lw=0.5, label=name)
        ax2.plot(df['Ship Speed (knots)'], df['Final Predicted CO2 Emission XGB After MC (kg)'], lw=0.5, label=name)
    fc_all = [combined_data[combined_data['Ship Speed (knots)']==s]['Final Predicted Fuel Consumption XGB After MC (kg/h)'].mean() for s in speeds]
    co2_all= [combined_data[combined_data['Ship Speed (knots)']==s]['Final Predicted CO2 Emission XGB After MC (kg)'].mean() for s in speeds]
    ax1.plot(speeds, fc_all, color='black', lw=1, label='Combined')
    ax2.plot(speeds, co2_all, color='black', lw=1, label='Combined')
    ax1.scatter(user_speed, fc_pred, color='red', zorder=5, label='You')
    ax2.scatter(user_speed, co2_pred,color='red', zorder=5, label='You')
    for ax in (ax1,ax2):
        ax.set_xlabel('Speed (knots)'); ax.grid(True); ax.legend(fontsize=6); ax.tick_params(labelsize=8)
    ax1.set_ylabel('FC (kg/h)'); ax2.set_ylabel('CO\u2082 (kg)')
    st.pyplot(fig)

# --- Report Function ---
def make_report(fc_pred, co2_pred):
    dists = {}
    for name, df in excel_data.items():
        fc_scen = df['Final Predicted Fuel Consumption XGB After MC (kg/h)'].mean()
        co2_scen= df['Final Predicted CO2 Emission XGB After MC (kg)'].mean()
        dists[name] = np.hypot(fc_pred-fc_scen, co2_pred-co2_scen)
    closest = sorted(dists, key=dists.get)
    return f"**Closest:** {closest[0]}\n\n**2nd closest:** {closest[1]}\n"

# --- UI Tabs ---
tab1, tab2, tab3 = st.tabs(["Predict", "Report", "Scenarios"])

with tab1:
    st.header("Step 1: Enter inputs")
    inputs = {}
    for feat,(mn,mx) in RANGES.items():
        step = (mx - mn)/100.0
        inputs[feat] = st.number_input(f"{feat}", min_value=mn, max_value=mx, value=(mn+mx)/2.0, step=step, format="%0.3f")

    if st.button("Run prediction"):
        fc_pred, co2_pred = predict_from_inputs(inputs)
        st.metric("Predicted FC (kg/h)", f"{fc_pred:.2f}")
        st.metric("Predicted CO\u2082 (kg)", f"{co2_pred:.2f}")
        make_plots(inputs['Ship Speed (knots)'], fc_pred, co2_pred)
        st.session_state["last"] = (inputs, fc_pred, co2_pred)

with tab2:
    st.header("Step 2: Report")
    if "last" not in st.session_state:
        st.info("\u25b6 Run a prediction in the Predict tab first.")
    else:
        _, fc_pred, co2_pred = st.session_state["last"]
        report_md = make_report(fc_pred, co2_pred)
        st.markdown(report_md)
        if st.button("Save report to Word"):
            doc = Document()
            doc.add_heading("Prediction Report", level=1)
            doc.add_paragraph(report_md)
            doc.save(report_path)
            st.success(f"Saved to {report_path}")

with tab3:
    st.header("Browse Scenario Data")
    sel = st.selectbox("Choose scenario", sheet_names)
    df = excel_data[sel][[
        'Ship Speed (knots)', 'Pe (kW)', 'Hull Efficiency (\u03b7H)',
        'Propeller Open Water Efficiency (\u03b7O)',
        'Final Predicted Fuel Consumption XGB After MC (kg/h)',
        'Final Predicted CO2 Emission XGB After MC (kg)'
    ]]
    st.dataframe(df, use_container_width=True)
