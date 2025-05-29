import tkinter as tk
from tkinter import ttk, messagebox
import pandas as pd
import numpy as np
import joblib
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from docx import Document
import os
from datetime import datetime

# File paths
model_path = "/Users/hamid/Documents/CO2 PAPER FINAL/PHASE_3_graphs/XGB_ML_model/"
data_path = "/Users/hamid/Documents/CO2 PAPER FINAL/PHASE_3_graphs/ML_Apply copy.xlsx"
output_path = "/Users/hamid/Documents/CO2 PAPER FINAL/PHASE_4_graphs/ML_Prediction_Results_1.xlsx"
report_path = "/Users/hamid/Documents/CO2 PAPER FINAL/PHASE_4_graphs/Prediction_Report.docx"

# Ensure output directory exists
os.makedirs(os.path.dirname(output_path), exist_ok=True)

# Feature ranges
RANGES = {
    'Pe (kW)': (300, 2200),
    'Ship Speed (knots)': (9, 15),
    'Hull Efficiency (ηH)': (1.0, 1.1),
    'Propeller Open Water Efficiency (ηO)': (0.5, 0.65)
}

# Load combined scenario data
try:
    sheet_names = [
        "Original Scenar_Coeff", "Paint (5%) Scen_Coeff", "Advance Propell_Coeff",
        "Fin (2%-4%) Sce_Coeff", "Bulbous Bow Sce_Coeff"
    ]
    excel_data = pd.read_excel(data_path, sheet_name=sheet_names)
    combined_data = pd.concat([excel_data[sheet] for sheet in sheet_names], ignore_index=True)
except Exception as e:
    print(f"Error reading Excel file: {e}")
    exit()

class PredictionApp:
    def __init__(self, root):
        self.root = root
        self.root.title("XGB Emission Prediction")
        self.root.geometry("1000x600")

        # Configure style for font
        style = ttk.Style()
        style.configure("TLabel", font=("Helvetica", 14))
        style.configure("TButton", font=("Helvetica", 12))

        # Main frame
        self.main_frame = ttk.Frame(root, padding="5")
        self.main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # Input fields
        self.entries = {}
        row = 0
        for feature, (min_val, max_val) in RANGES.items():
            label = f"{feature} ({min_val}–{max_val}):"
            ttk.Label(self.main_frame, text=label).grid(row=row, column=0, sticky=tk.W, pady=2)
            entry = ttk.Entry(self.main_frame, width=20)
            entry.grid(row=row, column=1, sticky=tk.W, pady=2)
            self.entries[feature] = entry
            row += 1

        # Buttons
        button_frame = ttk.Frame(self.main_frame)
        button_frame.grid(row=row, column=0, columnspan=2, pady=5)
        self.predict_button = ttk.Button(button_frame, text="Step 1: Predict", command=self.predict)
        self.predict_button.grid(row=0, column=0, padx=3)
        self.save_button = ttk.Button(button_frame, text="Save Results", command=self.save_results, state='disabled')
        self.save_button.grid(row=0, column=1, padx=3)
        self.report_button = ttk.Button(button_frame, text="Step 2: Generate Report", command=self.generate_report, state='disabled')
        self.report_button.grid(row=0, column=2, padx=3)
        self.zoom_in_button = ttk.Button(button_frame, text="+ Zoom In", command=self.zoom_in)
        self.zoom_in_button.grid(row=0, column=3, padx=3)
        self.zoom_out_button = ttk.Button(button_frame, text="- Zoom Out", command=self.zoom_out)
        self.zoom_out_button.grid(row=0, column=4, padx=3)

        # Result labels
        self.fc_label = ttk.Label(self.main_frame, text="Predicted FC (After MC): N/A")
        self.fc_label.grid(row=row+1, column=0, columnspan=2, pady=2)
        self.co2_label = ttk.Label(self.main_frame, text="Predicted CO2 (After MC): N/A")
        self.co2_label.grid(row=row+2, column=0, columnspan=2, pady=2)

        # Report display
        self.report_text = tk.Text(self.main_frame, height=5, width=80, font=("Helvetica", 15))
        self.report_text.grid(row=row+3, column=0, columnspan=2, pady=5)
        self.report_text.insert(tk.END, "Prediction report will appear here after Step 2.")
        self.report_text.config(state='disabled')

        # Matplotlib figure with two subplots
        self.base_figsize = (6, 6)  # Balanced aspect ratio
        self.figsize_scale = 1.0
        self.fig = plt.Figure(figsize=self.base_figsize)
        self.ax1 = self.fig.add_subplot(211)
        self.ax2 = self.fig.add_subplot(212)
        self.canvas = FigureCanvasTkAgg(self.fig, master=self.main_frame)
        self.canvas.get_tk_widget().grid(row=row+4, column=0, columnspan=2, pady=5)

        # Initialize variables
        self.xgb_model_fc = None
        self.xgb_model_co2 = None
        self.scaler = None
        self.last_prediction = None
        self.user_predictions = []

        # Load combined model
        try:
            self.xgb_model_fc = joblib.load(f"{model_path}xgb_model_fc_combined.pkl")
            self.xgb_model_co2 = joblib.load(f"{model_path}xgb_model_co2_combined.pkl")
            self.scaler = joblib.load(f"{model_path}scaler_combined.pkl")
            print("Combined model loaded successfully.")
        except FileNotFoundError:
            messagebox.showerror("Error", "Combined model files not found.")
            self.xgb_model_fc = None
            self.xgb_model_co2 = None
            self.scaler = None

    def zoom_in(self):
        self.figsize_scale *= 1.1
        self.update_figure_size()

    def zoom_out(self):
        self.figsize_scale = max(0.5, self.figsize_scale * 0.9)
        self.update_figure_size()

    def update_figure_size(self):
        # Destroy existing canvas
        self.canvas.get_tk_widget().destroy()
        # Create new figure
        new_figsize = (self.base_figsize[0] * self.figsize_scale, self.base_figsize[1] * self.figsize_scale)
        self.fig = plt.Figure(figsize=new_figsize)
        self.ax1 = self.fig.add_subplot(211)
        self.ax2 = self.fig.add_subplot(212)
        self.canvas = FigureCanvasTkAgg(self.fig, master=self.main_frame)
        self.canvas.get_tk_widget().grid(row=8, column=0, columnspan=2, pady=5)
        if self.last_prediction:
            self.predict()
        self.canvas.draw()

    def predict(self):
        if not self.xgb_model_fc or not self.xgb_model_co2 or not self.scaler:
            messagebox.showerror("Error", "Model not loaded.")
            return

        # Validate inputs
        inputs = {}
        for feature, entry in self.entries.items():
            try:
                value = float(entry.get())
                min_val, max_val = RANGES[feature]
                if not (min_val <= value <= max_val):
                    messagebox.showerror("Error", f"{feature} must be between {min_val} and {max_val}.")
                    return
                inputs[feature] = value
            except ValueError:
                messagebox.showerror("Error", f"Invalid input for {feature}.")
                return

        # Estimate remaining features using k-NN (k=3)
        speed = inputs['Ship Speed (knots)']
        pe = inputs['Pe (kW)']
        distances = np.sqrt(
            ((combined_data['Pe (kW)'] - pe) / RANGES['Pe (kW)'][1])**2 +
            ((combined_data['Ship Speed (knots)'] - speed) / RANGES['Ship Speed (knots)'][1])**2
        )
        nearest_idx = distances.nsmallest(3).index
        nearest_rows = combined_data.loc[nearest_idx]
        weights = 1 / (distances.loc[nearest_idx] + 1e-6)
        weights /= weights.sum()
        n_rpm = (nearest_rows['n (rpm)'] * weights).sum()
        t_kn = (nearest_rows['T (kN)'] * weights).sum()
        pd_kw = pe
        pb_kw = pe * 0.95

        # Prepare input data
        user_data = pd.DataFrame({
            'Pe (kW)': [pe],
            'n (rpm)': [n_rpm],
            'T (kN)': [t_kn],
            'Hull Efficiency (ηH)': [inputs['Hull Efficiency (ηH)']],
            'Propeller Open Water Efficiency (ηO)': [inputs['Propeller Open Water Efficiency (ηO)']],
            'Delivered Power (Pd) (kW)': [pd_kw],
            'Brake Power (PB) (kW)': [pb_kw],
            'Ship Speed (knots)': [speed],
            'Physics-Based FC': [pd_kw * inputs['Hull Efficiency (ηH)'] * 0.8],
            'Physics-Based CO2': [pd_kw * inputs['Propeller Open Water Efficiency (ηO)'] * 0.5],
            'Ship Speed^2': [speed ** 2],
            'Ship Speed^3': [speed ** 3],
            'Pd * Hull Efficiency': [pd_kw * inputs['Hull Efficiency (ηH)']],
            'Propeller Efficiency * n': [inputs['Propeller Open Water Efficiency (ηO)'] * n_rpm]
        })

        # Scale and predict
        user_data_scaled = self.scaler.transform(user_data)
        pred_fc = self.xgb_model_fc.predict(user_data_scaled)[0]
        pred_co2 = self.xgb_model_co2.predict(user_data_scaled)[0]

        # Update result labels
        self.fc_label.config(text=f"Predicted FC (After MC): {pred_fc:.2f} kg/h")
        self.co2_label.config(text=f"Predicted CO2 (After MC): {pred_co2:.2f} kg")

        # Store prediction
        self.last_prediction = {
            'Pe (kW)': pe,
            'Ship Speed (knots)': speed,
            'Hull Efficiency (ηH)': inputs['Hull Efficiency (ηH)'],
            'Propeller Open Water Efficiency (ηO)': inputs['Propeller Open Water Efficiency (ηO)'],
            'Predicted FC (After MC) (kg/h)': pred_fc,
            'Predicted CO2 (After MC) (kg)': pred_co2,
            'Type': 'User Prediction'
        }
        self.user_predictions.append(self.last_prediction)

        # Enable save button
        self.save_button.config(state='normal')

        # Update line graphs with prediction bar
        speeds = sorted(combined_data['Ship Speed (knots)'].unique())
        fc_values = []
        co2_values = []
        for s in speeds:
            subset = combined_data[combined_data['Ship Speed (knots)'] == s]
            fc_values.append(subset['Final Predicted Fuel Consumption XGB After MC (kg/h)'].mean())
            co2_values.append(subset['Final Predicted CO2 Emission XGB After MC (kg)'].mean())

        self.ax1.clear()
        self.ax2.clear()

        self.ax1.plot(speeds, fc_values, color='#1f77b4', label='Combined FC (After MC)')
        self.ax1.bar([speed], [pred_fc], color='#4d4d4d', width=0.2, label='User Prediction')
        self.ax1.set_title('Fuel Consumption (After MC)', fontsize=9)
        self.ax1.set_xlabel('Ship Speed (knots)', fontsize=9)
        self.ax1.set_ylabel('FC (kg/h)', fontsize=9)
        self.ax1.legend(fontsize=4, loc='upper left', frameon=True)
        self.ax1.grid(True)
        self.ax1.tick_params(axis='both', labelsize=8)
        self.ax1.legend().get_frame().set_linewidth(0.5)

        self.ax2.plot(speeds, co2_values, color='#1f77b4', label='Combined CO2 (After MC)')
        self.ax2.bar([speed], [pred_co2], color='#4d4d4d', width=0.2, label='User Prediction')
        self.ax2.set_title('CO2 Emissions (After MC)', fontsize=9)
        self.ax2.set_xlabel('Ship Speed (knots)', fontsize=9)
        self.ax2.set_ylabel('CO2 (kg)', fontsize=9)
        self.ax2.legend(fontsize=4, loc='upper left', frameon=True)
        self.ax2.grid(True)
        self.ax2.tick_params(axis='both', labelsize=8)
        self.ax2.legend().get_frame().set_linewidth(0.5)

        self.fig.tight_layout()
        self.canvas.draw()
        self.canvas.flush_events()

    def save_results(self):
        if not self.last_prediction:
            messagebox.showerror("Error", "No prediction to save.")
            return

        user_data = pd.DataFrame(self.user_predictions)
        model_data = combined_data[[
            'Ship Speed (knots)', 'Final Predicted Fuel Consumption XGB After MC (kg/h)',
            'Final Predicted CO2 Emission XGB After MC (kg)'
        ]].copy()
        model_data['Type'] = 'ML Model (After MC)'

        try:
            with pd.ExcelWriter(output_path, engine='openpyxl', mode='a', if_sheet_exists='replace') as writer:
                user_data.to_excel(writer, sheet_name='User_Predictions', index=False)
                model_data.to_excel(writer, sheet_name='Model_Predictions', index=False)
            messagebox.showinfo("Success", f"Results saved to {output_path}")
            self.save_button.config(state='disabled')
            self.report_button.config(state='normal')
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save results: {e}")

    def generate_report(self):
        if not self.user_predictions:
            messagebox.showerror("Error", "No saved predictions to generate report.")
            return

        latest_pred = self.user_predictions[-1]
        speed = latest_pred['Ship Speed (knots)']
        pred_fc = latest_pred['Predicted FC (After MC) (kg/h)']
        pred_co2 = latest_pred['Predicted CO2 (After MC) (kg)']

        # Compare to scenarios
        scenario_distances = {}
        for sheet_name in sheet_names:
            scenario_data = excel_data[sheet_name]
            if speed in scenario_data['Ship Speed (knots)'].values:
                row = scenario_data[scenario_data['Ship Speed (knots)'] == speed].iloc[0]
                fc = row['Final Predicted Fuel Consumption XGB After MC (kg/h)']
                co2 = row['Final Predicted CO2 Emission XGB After MC (kg)']
            else:
                closest_speed = scenario_data['Ship Speed (knots)'].iloc[
                    (scenario_data['Ship Speed (knots)'] - speed).abs().idxmin()
                ]
                row = scenario_data[scenario_data['Ship Speed (knots)'] == closest_speed].iloc[0]
                fc = row['Final Predicted Fuel Consumption XGB After MC (kg/h)']
                co2 = row['Final Predicted CO2 Emission XGB After MC (kg)']
            distance = np.sqrt((pred_fc - fc)**2 + (pred_co2 - co2)**2)
            scenario_distances[sheet_name] = {'distance': distance, 'fc': fc, 'co2': co2}

        # Find closest scenarios
        sorted_scenarios = sorted(scenario_distances.items(), key=lambda x: x[1]['distance'])
        closest_scenario = sorted_scenarios[0][0]
        second_closest = sorted_scenarios[1][0] if len(sorted_scenarios) > 1 else None

        # Suggest improvement
        min_fc_scenario = min(
            scenario_distances.items(),
            key=lambda x: x[1]['fc']
        )[0]
        suggestion = f"For lower FC and CO₂, consider improvements similar to {min_fc_scenario}."

        # Generate report text
        report_text = (
            f"Prediction Report (Generated {datetime.now().strftime('%Y-%m-%d %H:%M:%S')})\n"
            f"Input: Pe={latest_pred['Pe (kW)']:.2f} kW, Speed={speed:.1f} knots, "
            f"ηH={latest_pred['Hull Efficiency (ηH)']:.3f}, ηO={latest_pred['Propeller Open Water Efficiency (ηO)']:.3f}\n"
            f"Predicted FC: {pred_fc:.2f} kg/h, CO2: {pred_co2:.2f} kg\n"
            f"Closest Scenario: {closest_scenario} (FC={scenario_distances[closest_scenario]['fc']:.2f}, "
            f"CO2={scenario_distances[closest_scenario]['co2']:.2f})\n"
        )
        if second_closest:
            report_text += (
                f"Second Closest: {second_closest} (FC={scenario_distances[second_closest]['fc']:.2f}, "
                f"CO2={scenario_distances[second_closest]['co2']:.2f})\n"
            )
        report_text += suggestion

        # Display report
        self.report_text.config(state='normal')
        self.report_text.delete(1.0, tk.END)
        self.report_text.insert(tk.END, report_text)
        self.report_text.config(state='disabled')

        # Save report to Word
        doc = Document()
        doc.add_heading('Prediction Report', 0)
        doc.add_paragraph(report_text)
        try:
            doc.save(report_path)
            messagebox.showinfo("Success", f"Report saved to {report_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save report: {e}")

if __name__ == "__main__":
    root = tk.Tk()
    app = PredictionApp(root)
    root.mainloop()