import streamlit as st
import pandas as pd
import numpy as np
import joblib
import matplotlib.pyplot as plt
from docx import Document
import sys, os


# ---- Header and Branding ----
col1, col2 = st.columns([6, 1])
with col1:
    st.markdown("<h1 style='margin-bottom: 0;'>🚢 XGB CO₂ & Fuel-Consumption Predictor</h1>", unsafe_allow_html=True)
with col2:
    st.image("AppIcon.png", width=100)

col3, col4 = st.columns([6, 2])
with col4:
    st.markdown(
        "<div style='text-align: right; font-size: 0.9em;'>"
        "Created by Hamid<br>© 2025 All rights reserved"
        "</div>", unsafe_allow_html=True
    )

with st.expander("ℹ️ About this App"):
    st.markdown("""
    **Project Title:** CO₂ & Fuel Consumption Prediction using XGBoost  
    **Developed by:**  
    - Hamid, PhD Researcher in Maritime Energy & ML  
    **Acknowledgments:**  
    This tool was developed as part of a sustainability-driven maritime research project focused on CO₂ & Fuel-Consumption prediction models.
    """)
#–– allow bundled vs. dev paths
if getattr(sys, "frozen", False):
    base_path = sys._MEIPASS
else:
    base_path = os.path.dirname(__file__)

model_path  = os.path.join(base_path, "XGB_ML_model") + os.sep
data_path   = os.path.join(base_path, "ML_Apply copy.xlsx")
output_path = os.path.join(base_path, "ML_Prediction_Results_1.xlsx")
report_path = os.path.join(base_path, "Prediction_Report.docx")

# load your models & scaler
xgb_fc  = joblib.load(f"{model_path}xgb_model_fc_combined.pkl")
xgb_co2 = joblib.load(f"{model_path}xgb_model_co2_combined.pkl")
scaler  = joblib.load(f"{model_path}scaler_combined.pkl")

# feature ranges
RANGES = {
    'Pe (kW)': (300, 2200),
    'Ship Speed (knots)': (9, 15),
    'Hull Efficiency (ηH)': (1.0, 1.1),
    'Propeller Open Water Efficiency (ηO)': (0.5, 0.65)
}

FEATURE_COLS = [
    'Pe (kW)', 'n (rpm)', 'T (kN)',
    'Hull Efficiency (ηH)', 'Propeller Open Water Efficiency (ηO)',
    'Delivered Power (Pd) (kW)', 'Brake Power (PB) (kW)', 'Ship Speed (knots)',
    'Physics-Based FC','Physics-Based CO2',
    'Ship Speed^2','Ship Speed^3',
    'Pd * Hull Efficiency','Propeller Efficiency * n'
]

# load all sheets
sheet_names = [
    "Original Scenar_Coeff", "Paint (5%) Scen_Coeff", "Advance Propell_Coeff",
    "Fin (2%-4%) Sce_Coeff", "Bulbous Bow Sce_Coeff"
]
excel_data = pd.read_excel(data_path, sheet_name=sheet_names)
combined_data = pd.concat(excel_data.values(), ignore_index=True)

# helper: do kNN fill + predict
def predict_from_inputs(vals):
    df = combined_data
    dist = np.sqrt(
        ((df['Pe (kW)']-vals['Pe (kW)'])/RANGES['Pe (kW)'][1])**2 +
        ((df['Ship Speed (knots)']-vals['Ship Speed (knots)'])/RANGES['Ship Speed (knots)'][1])**2
    )
    idx = dist.nsmallest(3).index
    w = 1/(dist.loc[idx]+1e-6); w/=w.sum()
    nrpm = (df.loc[idx,'n (rpm)']*w).sum()
    tkn  = (df.loc[idx,'T (kN)']*w).sum()
    pd_kw, pb_kw = vals['Pe (kW)'], vals['Pe (kW)']*0.95

    row = {
        **vals,
        'n (rpm)': nrpm,
        'T (kN)':  tkn,
        'Delivered Power (Pd) (kW)': pd_kw,
        'Brake Power (PB) (kW)': pb_kw,
        'Physics-Based FC': pd_kw*vals['Hull Efficiency (ηH)']*0.8,
        'Physics-Based CO2': pd_kw*vals['Propeller Open Water Efficiency (ηO)']*0.5,
        'Ship Speed^2': vals['Ship Speed (knots)']**2,
        'Ship Speed^3': vals['Ship Speed (knots)']**3,
        'Pd * Hull Efficiency': pd_kw*vals['Hull Efficiency (ηH)'],
        'Propeller Efficiency * n': nrpm*vals['Propeller Open Water Efficiency (ηO)']
    }
    X = pd.DataFrame([row])[FEATURE_COLS]
    Xs = scaler.transform(X)
    return float(xgb_fc.predict(Xs)), float(xgb_co2.predict(Xs))

# helper: build combined + scenario plots
def make_plots(user_speed, fc_pred, co2_pred):
    fig, (ax1,ax2) = plt.subplots(2,1, figsize=(6,6))
    speeds = sorted(combined_data['Ship Speed (knots)'].unique())
    # scenario lines
    for name, df in excel_data.items():
        ax1.plot(df['Ship Speed (knots)'], df['Final Predicted Fuel Consumption XGB After MC (kg/h)'], lw=0.5, label=name)
        ax2.plot(df['Ship Speed (knots)'], df['Final Predicted CO2 Emission XGB After MC (kg)'], lw=0.5, label=name)
    # combined
    fc_all = [combined_data[combined_data['Ship Speed (knots)']==s]['Final Predicted Fuel Consumption XGB After MC (kg/h)'].mean() for s in speeds]
    co2_all= [combined_data[combined_data['Ship Speed (knots)']==s]['Final Predicted CO2 Emission XGB After MC (kg)'].mean() for s in speeds]
    ax1.plot(speeds, fc_all, color='black', lw=1, label='Combined')
    ax2.plot(speeds, co2_all, color='black', lw=1, label='Combined')
    # user point
    ax1.scatter(user_speed, fc_pred, color='red', zorder=5, label='You')
    ax2.scatter(user_speed, co2_pred,color='red', zorder=5, label='You')
    for ax in (ax1,ax2):
        ax.set_xlabel('Speed (knots)'); ax.grid(True); ax.legend(fontsize=6); ax.tick_params(labelsize=8)
    ax1.set_ylabel('FC (kg/h)'); ax2.set_ylabel('CO₂ (kg)')
    st.pyplot(fig)

# helper: make text report
def make_report(fc_pred, co2_pred):
    dists = {}
    for name, df in excel_data.items():
        fc_scen = df['Final Predicted Fuel Consumption XGB After MC (kg/h)'].mean()
        co2_scen= df['Final Predicted CO2 Emission XGB After MC (kg)'].mean()
        dists[name] = np.hypot(fc_pred-fc_scen, co2_pred-co2_scen)
    closest = sorted(dists, key=dists.get)
    return (
        f"**Closest:** {closest[0]}\n\n"
        f"**2nd closest:** {closest[1]}\n"
    )

st.title("🚢 XGB CO₂ & Fuel-Consumption Predictor")

tab1, tab2, tab3 = st.tabs(["Predict", "Report", "Scenarios"])

with tab1:
    st.header("Step 1: Enter inputs")
    inputs = {}
    for feat,(mn,mx) in RANGES.items():
        # force floats for all three
        mn_f, mx_f = float(mn), float(mx)
        default  = (mn_f + mx_f) / 2.0
        step     = (mx_f - mn_f) / 100.0
        inputs[feat] = st.number_input(
            f"{feat}", min_value=mn_f, max_value=mx_f,
            value=default, step=step
        )

    if st.button("Run prediction"):
        fc_pred, co2_pred = predict_from_inputs(inputs)
        st.metric("Predicted FC (kg/h)", f"{fc_pred:.2f}")
        st.metric("Predicted CO₂  (kg)", f"{co2_pred:.2f}")
        make_plots(inputs['Ship Speed (knots)'], fc_pred, co2_pred)
        st.session_state["last"] = (inputs, fc_pred, co2_pred)

with tab2:
    st.header("Step 2: Report")
    if "last" not in st.session_state:
        st.info("▶ Run a prediction in the Predict tab first.")
    else:
        _, fc_pred, co2_pred = st.session_state["last"]
        report_md = make_report(fc_pred, co2_pred)
        st.markdown(report_md)
        if st.button("Save report to Word"):
            doc = Document()
            doc.add_heading("Prediction Report", level=1)
            doc.add_paragraph(report_md)
            doc.save(report_path)
            st.success(f"Saved to {report_path}")

with tab3:
    st.header("Browse Scenario Data")
    sel = st.selectbox("Choose scenario", sheet_names)
    df = excel_data[sel][[
        'Ship Speed (knots)', 'Pe (kW)', 'Hull Efficiency (ηH)',
        'Propeller Open Water Efficiency (ηO)',
        'Final Predicted Fuel Consumption XGB After MC (kg/h)',
        'Final Predicted CO2 Emission XGB After MC (kg)'
    ]]
    st.dataframe(df, use_container_width=True)
