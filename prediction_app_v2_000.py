import tkinter as tk
from tkinter import ttk, messagebox
import pandas as pd
import numpy as np
import joblib
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from docx import Document
import os
from datetime import datetime

# File paths
model_path = "/Users/hamid/Documents/CO2 PAPER FINAL/PHASE_3_graphs/XGB_ML_model/"
data_path = "/Users/hamid/Documents/CO2 PAPER FINAL/PHASE_3_graphs/ML_Apply copy.xlsx"
output_path = "/Users/hamid/Documents/CO2 PAPER FINAL/PHASE_4_graphs/ML_Prediction_Results_1.xlsx"
report_path = "/Users/hamid/Documents/CO2 PAPER FINAL/PHASE_4_graphs/Prediction_Report.docx"

# Ensure output directory exists
os.makedirs(os.path.dirname(output_path), exist_ok=True)

# Feature ranges
RANGES = {
    'Pe (kW)': (300, 2200),
    'Ship Speed (knots)': (9, 15),
    'Hull Efficiency (ηH)': (1.0, 1.1),
    'Propeller Open Water Efficiency (ηO)': (0.5, 0.65)
}

# Columns order for scaler
FEATURE_COLS = [
    'Pe (kW)', 'n (rpm)', 'T (kN)',
    'Hull Efficiency (ηH)', 'Propeller Open Water Efficiency (ηO)',
    'Delivered Power (Pd) (kW)', 'Brake Power (PB) (kW)', 'Ship Speed (knots)',
    'Physics-Based FC', 'Physics-Based CO2',
    'Ship Speed^2', 'Ship Speed^3',
    'Pd * Hull Efficiency', 'Propeller Efficiency * n'
]

# Columns to show in scenario tab
DISPLAY_COLS = [
    'Ship Speed (knots)', 'Pe (kW)', 'Hull Efficiency (ηH)',
    'Propeller Open Water Efficiency (ηO)',
    'Final Predicted Fuel Consumption XGB After MC (kg/h)',
    'Final Predicted CO2 Emission XGB After MC (kg)'
]

# Load scenario data
sheet_names = [
    "Original Scenar_Coeff", "Paint (5%) Scen_Coeff", "Advance Propell_Coeff",
    "Fin (2%-4%) Sce_Coeff", "Bulbous Bow Sce_Coeff"
]
excel_data = pd.read_excel(data_path, sheet_name=sheet_names)
combined_data = pd.concat([excel_data[s] for s in sheet_names], ignore_index=True)

class PredictionApp:
    def __init__(self, root):
        self.root = root
        self.root.title("XGB Emission Prediction")
        self.root.geometry("1000x600")

        style = ttk.Style()
        style.configure("TLabel", font=("Helvetica", 14))
        style.configure("TButton", font=("Helvetica", 12))

        # Notebook
        nb = ttk.Notebook(root)
        nb.pack(fill='both', expand=True)

        # Prediction tab
        self.pred_frame = ttk.Frame(nb)
        nb.add(self.pred_frame, text='Prediction')
        self._build_prediction_tab()

        # Scenario tab
        data_frame = ttk.Frame(nb)
        nb.add(data_frame, text='Scenario Data')
        self._build_data_tab(data_frame)

    def _build_prediction_tab(self):
        f = self.pred_frame
        # inputs
        frm = ttk.Frame(f, padding=5)
        frm.pack(fill='x')
        self.entries = {}
        for i,(feat,(mn,mx)) in enumerate(RANGES.items()):
            ttk.Label(frm, text=f"{feat} ({mn}-{mx}):").grid(row=i,column=0,sticky='w',padx=5,pady=2)
            e = ttk.Entry(frm, width=15)
            e.grid(row=i,column=1,padx=5,pady=2)
            self.entries[feat]=e
        # buttons
        bf = ttk.Frame(frm)
        bf.grid(row=len(RANGES),column=0,columnspan=5,pady=5)
        self.b1=ttk.Button(bf,text='Step 1: Predict',command=self.predict)
        self.b1.grid(row=0,column=0,padx=3)
        self.b_save=ttk.Button(bf,text='Save Results',command=self.save_results,state='disabled')
        self.b_save.grid(row=0,column=1,padx=3)
        self.b2=ttk.Button(bf,text='Step 2: Generate Report',command=self.generate_report,state='disabled')
        self.b2.grid(row=0,column=2,padx=3)
        self.zoom_in_button=ttk.Button(bf,text='+ Zoom In',command=self.zoom_in)
        self.zoom_in_button.grid(row=0,column=3,padx=3)
        self.zoom_out_button=ttk.Button(bf,text='- Zoom Out',command=self.zoom_out)
        self.zoom_out_button.grid(row=0,column=4,padx=3)
        # results
        self.fc_lbl=ttk.Label(f,text='Predicted FC: N/A')
        self.fc_lbl.pack(anchor='w',padx=10)
        self.co2_lbl=ttk.Label(f,text='Predicted CO2: N/A')
        self.co2_lbl.pack(anchor='w',padx=10)
        # report
        self.report_text=tk.Text(f,height=5,width=80)
        self.report_text.pack(padx=10,pady=5)
        self.report_text.insert('end','Report will appear here.')
        self.report_text.config(state='disabled')
        # plot
        self.base_figsize=(6,6)
        self.fig=plt.Figure(figsize=self.base_figsize)
        self.ax1=self.fig.add_subplot(211)
        self.ax2=self.fig.add_subplot(212)
        self.canvas=FigureCanvasTkAgg(self.fig,master=f)
        self.canvas.get_tk_widget().pack(fill='both',expand=True)
        self.figsize_scale=1.0
        # load models
        try:
            self.xgb_fc=joblib.load(f"{model_path}xgb_model_fc_combined.pkl")
            self.xgb_co2=joblib.load(f"{model_path}xgb_model_co2_combined.pkl")
            self.scaler=joblib.load(f"{model_path}scaler_combined.pkl")
        except Exception:
            messagebox.showerror('Error','Could not load models')
            root.destroy()
        self.last=None

    def zoom_in(self):
        self.figsize_scale*=1.1
        self._update_fig()
    def zoom_out(self):
        self.figsize_scale=max(0.5,self.figsize_scale*0.9)
        self._update_fig()

    def _update_fig(self):
        self.canvas.get_tk_widget().destroy()
        sz=(self.base_figsize[0]*self.figsize_scale, self.base_figsize[1]*self.figsize_scale)
        self.fig=plt.Figure(figsize=sz)
        self.ax1=self.fig.add_subplot(211)
        self.ax2=self.fig.add_subplot(212)
        self.canvas=FigureCanvasTkAgg(self.fig, master=self.pred_frame)
        self.canvas.get_tk_widget().pack(fill='both',expand=True)
        if self.last:
            self._plot_results(*self.last)
        self.canvas.draw()

    def _build_data_tab(self, df):
        ttk.Label(df,text='Select Scenario:').pack(anchor='w',padx=5,pady=5)
        self.sc_var=tk.StringVar()
        dd=ttk.Combobox(df,textvariable=self.sc_var,values=sheet_names,state='readonly')
        dd.pack(anchor='w',padx=5)
        dd.bind('<<ComboboxSelected>>',self._show_table)
        self.tv=ttk.Treeview(df,show='headings')
        self.tv.pack(fill='both',expand=True,padx=5,pady=5)

    def _show_table(self,e):
        name=self.sc_var.get()
        data=excel_data[name][DISPLAY_COLS]
        self.tv.delete(*self.tv.get_children())
        self.tv['columns']=DISPLAY_COLS
        for c in DISPLAY_COLS:
            self.tv.heading(c,text=c)
            self.tv.column(c,width=100)
        for _,row in data.iterrows():
            self.tv.insert('', 'end', values=list(row))

    def predict(self):
        vals={}
        for k,e in self.entries.items():
            try: v=float(e.get())
            except: return messagebox.showerror('Error',f'Invalid {k}')
            mn,mx=RANGES[k]
            if not(mn<=v<=mx): return messagebox.showerror('Error',f'{k} out of range')
            vals[k]=v
        df=combined_data
        dist=np.sqrt(((df['Pe (kW)']-vals['Pe (kW)'])/RANGES['Pe (kW)'][1])**2+((df['Ship Speed (knots)']-vals['Ship Speed (knots)'])/RANGES['Ship Speed (knots)'][1])**2)
        idx=dist.nsmallest(3).index
        w=1/(dist.loc[idx]+1e-6);w/=w.sum()
        nrpm=(df.loc[idx,'n (rpm)']*w).sum(); tkn=(df.loc[idx,'T (kN)']*w).sum()
        pd_kw=vals['Pe (kW)']; pb_kw=pd_kw*0.95
        # build row with correct order
        row={**vals,'n (rpm)':nrpm,'T (kN)':tkn,'Delivered Power (Pd) (kW)':pd_kw,'Brake Power (PB) (kW)':pb_kw,
             'Physics-Based FC':pd_kw*vals['Hull Efficiency (ηH)']*0.8,'Physics-Based CO2':pd_kw*vals['Propeller Open Water Efficiency (ηO)']*0.5,
             'Ship Speed^2':vals['Ship Speed (knots)']**2,'Ship Speed^3':vals['Ship Speed (knots)']**3,
             'Pd * Hull Efficiency':pd_kw*vals['Hull Efficiency (ηH)'],'Propeller Efficiency * n':nrpm*vals['Propeller Open Water Efficiency (ηO)']}
        X=pd.DataFrame([row])[FEATURE_COLS]
        Xs=self.scaler.transform(X)
        fcp=self.xgb_fc.predict(Xs)[0]; co2p=self.xgb_co2.predict(Xs)[0]
        self.fc_lbl.config(text=f"Predicted FC: {fcp:.2f} kg/h")
        self.co2_lbl.config(text=f"Predicted CO2: {co2p:.2f} kg")
        self.last=(vals,fcp,co2p)
        self.b_save.config(state='normal'); self.b2.config(state='normal')
        self._plot_results(vals,fcp,co2p)

    def _plot_results(self, vals, fcp, co2p):
        df=combined_data
        speeds=sorted(df['Ship Speed (knots)'].unique())
        self.ax1.clear(); self.ax2.clear()
        # scenario lines
        for name in sheet_names:
            sd=excel_data[name]
            self.ax1.plot(sd['Ship Speed (knots)'],sd['Final Predicted Fuel Consumption XGB After MC (kg/h)'],lw=0.5,label=name)
            self.ax2.plot(sd['Ship Speed (knots)'],sd['Final Predicted CO2 Emission XGB After MC (kg)'],lw=0.5,label=name)
        # combined
        fc_all=[df[df['Ship Speed (knots)']==s]['Final Predicted Fuel Consumption XGB After MC (kg/h)'].mean() for s in speeds]
        co_all=[df[df['Ship Speed (knots)']==s]['Final Predicted CO2 Emission XGB After MC (kg)'].mean() for s in speeds]
        self.ax1.plot(speeds,fc_all,color='black',lw=1,label='Combined')
        self.ax2.plot(speeds,co_all,color='black',lw=1,label='Combined')
        # user
        v=vals['Ship Speed (knots)']
        self.ax1.scatter(v,fcp,color='red',zorder=5,label='User')
        self.ax2.scatter(v,co2p,color='red',zorder=5,label='User')
        for ax in (self.ax1,self.ax2):
            ax.tick_params(labelsize=8)
            ax.legend(fontsize=6)
        self.canvas.draw()

    def save_results(self):
        if not self.last: return
        vals,fcp,co2p=self.last
        df=pd.DataFrame([{**vals,'FC After MC':fcp,'CO2 After MC':co2p}])
        try:
            with pd.ExcelWriter(output_path,engine='openpyxl',mode='a',if_sheet_exists='replace') as w:
                df.to_excel(w,sheet_name='User_Predictions',index=False)
            messagebox.showinfo('Saved',f'Results saved to {output_path}')
            self.b_save.config(state='disabled')
        except Exception as e:
            messagebox.showerror('Error',str(e))

    def generate_report(self):
        if not self.last: return
        vals,fcp,co2p=self.last
        d={name:np.hypot(fcp-excel_data[name]['Final Predicted Fuel Consumption XGB After MC (kg/h)'].mean(),
                        co2p-excel_data[name]['Final Predicted CO2 Emission XGB After MC (kg)'].mean()) for name in sheet_names}
        s=sorted(d,key=d.get)
        text=(f"Input: {vals}\nPred FC: {fcp:.2f}, CO2: {co2p:.2f}\n"
              f"Closest: {s[0]}\nSecond Closest: {s[1]}\n")
        self.report_text.config(state='normal')
        self.report_text.delete('1.0','end')
        self.report_text.insert('end',text)
        self.report_text.config(state='disabled')
        doc=Document();doc.add_heading('Prediction Report',0);doc.add_paragraph(text)
        doc.save(report_path);messagebox.showinfo('Saved',f'Report saved to {report_path}')

if __name__=='__main__':
    root=tk.Tk()
    PredictionApp(root)
    root.mainloop()