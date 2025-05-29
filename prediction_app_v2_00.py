import tkinter as tk
from tkinter import ttk, messagebox
import pandas as pd
import numpy as np
import joblib
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from docx import Document
import os
from datetime import datetime

# File paths (adjust as needed)
model_path = "/Users/hamid/Documents/CO2 PAPER FINAL/PHASE_3_graphs/XGB_ML_model/"
data_path = "/Users/hamid/Documents/CO2 PAPER FINAL/PHASE_3_graphs/ML_Apply copy.xlsx"
output_path = "/Users/hamid/Documents/CO2 PAPER FINAL/PHASE_4_graphs/ML_Prediction_Results_1.xlsx"
report_path = "/Users/hamid/Documents/CO2 PAPER FINAL/PHASE_4_graphs/Prediction_Report.docx"

# Ensure output directory exists
os.makedirs(os.path.dirname(output_path), exist_ok=True)

# Feature ranges
RANGES = {
    'Pe (kW)': (300, 2200),
    'Ship Speed (knots)': (9, 15),
    'Hull Efficiency (ηH)': (1.0, 1.1),
    'Propeller Open Water Efficiency (ηO)': (0.5, 0.65)
}

# Columns to display in scenario tab
DISPLAY_COLS = [
    'Ship Speed (knots)', 'Pe (kW)', 'Hull Efficiency (ηH)',
    'Propeller Open Water Efficiency (ηO)',
    'Fuel Consumption (FC) (kg/h)', 'CO2 Emission (kg)',
    'Final Predicted Fuel Consumption XGB Before MC (kg/h)',
    'Final Predicted CO2 Emission XGB Before MC (kg)',
    'Final Predicted Fuel Consumption XGB After MC (kg/h)',
    'Final Predicted CO2 Emission XGB After MC (kg)'
]

# Load combined scenario data
sheet_names = [
    "Original Scenar_Coeff", "Paint (5%) Scen_Coeff", "Advance Propell_Coeff",
    "Fin (2%-4%) Sce_Coeff", "Bulbous Bow Sce_Coeff"
]
excel_data = pd.read_excel(data_path, sheet_name=sheet_names)
combined_data = pd.concat([excel_data[s] for s in sheet_names], ignore_index=True)


class PredictionApp:
    def __init__(self, root):
        self.root = root
        self.root.title("XGB Emission Prediction")
        self.root.geometry("1000x650")

        # Notebook for tabs
        self.nb = ttk.Notebook(root)
        self.nb.pack(fill='both', expand=True)

        # Prediction tab
        self.pred_frame = ttk.Frame(self.nb)
        self.nb.add(self.pred_frame, text='Prediction')

        # Scenario Data tab
        self.data_frame = ttk.Frame(self.nb)
        self.nb.add(self.data_frame, text='Scenario Data')
        self._build_data_tab()

        # Build prediction UI
        self._build_prediction_tab()

    def _build_prediction_tab(self):
        frame = self.pred_frame
        # Inputs
        inputs_frame = ttk.Frame(frame, padding=5)
        inputs_frame.pack(side='top', fill='x')
        self.entries = {}
        for i, (feat, (mn, mx)) in enumerate(RANGES.items()):
            ttk.Label(inputs_frame, text=f"{feat} ({mn}-{mx}):").grid(row=i, column=0, sticky='w', padx=5, pady=2)
            e = ttk.Entry(inputs_frame, width=15)
            e.grid(row=i, column=1, padx=5, pady=2)
            self.entries[feat] = e
        # Buttons
        btn_frame = ttk.Frame(inputs_frame)
        btn_frame.grid(row=len(RANGES), column=0, columnspan=5, pady=5)
        self.b1 = ttk.Button(btn_frame, text='Step 1: Predict', command=self.predict)
        self.b1.grid(row=0, column=0, padx=3)
        self.b_save = ttk.Button(btn_frame, text='Save Results', command=self.save_results, state='disabled')
        self.b_save.grid(row=0, column=1, padx=3)
        self.b2 = ttk.Button(btn_frame, text='Step 2: Generate Report', command=self.generate_report, state='disabled')
        self.b2.grid(row=0, column=2, padx=3)
        self.zin = ttk.Button(btn_frame, text='+ Zoom In', command=lambda: self.zoom(1.1))
        self.zin.grid(row=0, column=3, padx=3)
        self.zout = ttk.Button(btn_frame, text='- Zoom Out', command=lambda: self.zoom(0.9))
        self.zout.grid(row=0, column=4, padx=3)
        # Results labels
        self.fc_lbl = ttk.Label(frame, text='Predicted FC: N/A')
        self.fc_lbl.pack(anchor='w', padx=10)
        self.co2_lbl = ttk.Label(frame, text='Predicted CO2: N/A')
        self.co2_lbl.pack(anchor='w', padx=10)
        # Report box
        self.report_text = tk.Text(frame, height=6, width=100)
        self.report_text.pack(padx=10, pady=5)
        self.report_text.insert('end', 'Report will appear here.')
        self.report_text.config(state='disabled')
        # Plot area
        self.fig, (self.ax1, self.ax2) = plt.subplots(2, 1, figsize=(6, 6))
        self.canvas = FigureCanvasTkAgg(self.fig, master=frame)
        self.canvas.get_tk_widget().pack(fill='both', expand=True)
        self.scale = 1.0
        # load models
        try:
            self.xgb_fc = joblib.load(f"{model_path}xgb_model_fc_combined.pkl")
            self.xgb_co2 = joblib.load(f"{model_path}xgb_model_co2_combined.pkl")
            self.scaler = joblib.load(f"{model_path}scaler_combined.pkl")
        except Exception:
            messagebox.showerror('Error', 'Could not load models')
            root.destroy()
        self.last = None

    def _build_data_tab(self):
        df = self.data_frame
        ttk.Label(df, text='Select Scenario:').pack(anchor='w', padx=5, pady=5)
        self.sc_var = tk.StringVar()
        dd = ttk.Combobox(df, textvariable=self.sc_var, values=sheet_names, state='readonly')
        dd.pack(anchor='w', padx=5)
        dd.bind('<<ComboboxSelected>>', self.show_table)
        self.tv = ttk.Treeview(df, show='headings')
        self.tv.pack(fill='both', expand=True, padx=5, pady=5)

    def show_table(self, e):
        name = self.sc_var.get()
        data = excel_data[name][DISPLAY_COLS]
        self.tv.delete(*self.tv.get_children())
        self.tv['columns'] = DISPLAY_COLS
        for c in DISPLAY_COLS:
            self.tv.heading(c, text=c)
            self.tv.column(c, width=100)
        for _, row in data.iterrows():
            self.tv.insert('', 'end', values=list(row))

    def predict(self):
        vals = {}
        # validate inputs
        for k, e in self.entries.items():
            try:
                v = float(e.get())
            except:
                messagebox.showerror('Error', f'Invalid {k}')
                return
            mn, mx = RANGES[k]
            if not (mn <= v <= mx):
                messagebox.showerror('Error', f'{k} out of range')
                return
            vals[k] = v
        # k-NN for missing
        df = combined_data
        dist = np.sqrt(((df['Pe (kW)'] - vals['Pe (kW)']) / (RANGES['Pe (kW)'][1])) ** 2 +
                       ((df['Ship Speed (knots)'] - vals['Ship Speed (knots)']) / (
                       RANGES['Ship Speed (knots)'][1])) ** 2)
        idx = dist.nsmallest(3).index
        w = 1 / (dist.loc[idx] + 1e-6);
        w /= w.sum()
        nrpm = (df.loc[idx, 'n (rpm)'] * w).sum()
        tkn = (df.loc[idx, 'T (kN)'] * w).sum()
        pd_kw, pb_kw = vals['Pe (kW)'], vals['Pe (kW)'] * 0.95
        # build feature row
        row = {**vals,
               'n (rpm)': nrpm,
               'T (kN)': tkn,
               'Delivered Power (Pd) (kW)': pd_kw,
               'Brake Power (PB) (kW)': pb_kw,
               'Physics-Based FC': pd_kw * vals['Hull Efficiency (ηH)'] * 0.8,
               'Physics-Based CO2': pd_kw * vals['Propeller Open Water Efficiency (ηO)'] * 0.5,
               'Ship Speed^2': vals['Ship Speed (knots)'] ** 2,
               'Ship Speed^3': vals['Ship Speed (knots)'] ** 3,
               'Pd * Hull Efficiency': pd_kw * vals['Hull Efficiency (ηH)'],
               'Propeller Efficiency * n': nrpm * vals['Propeller Open Water Efficiency (ηO)']
               }
        X = pd.DataFrame([row])
        # ensure correct column order
        X = X[self.scaler.feature_names_in_]
        Xs = self.scaler.transform(X)
        fcp = self.xgb_fc.predict(Xs)[0]
        co2p = self.xgb_co2.predict(Xs)[0]
        # update labels
        self.fc_lbl.config(text=f"Predicted FC: {fcp:.2f} kg/h")
        self.co2_lbl.config(text=f"Predicted CO2: {co2p:.2f} kg")
        self.last = (vals, fcp, co2p)
        self.b_save.config(state='normal')
        # redraw plots
        speeds = sorted(df['Ship Speed (knots)'].unique())
        fc_all = [df[df['Ship Speed (knots)'] == s]['Final Predicted Fuel Consumption XGB After MC (kg/h)'].mean() for s
                  in speeds]
        co_all = [df[df['Ship Speed (knots)'] == s]['Final Predicted CO2 Emission XGB After MC (kg)'].mean() for s in
                  speeds]
        self.ax1.clear();
        self.ax2.clear()
        for name in sheet_names:
            sd = excel_data[name]
            self.ax1.plot(sd['Ship Speed (knots)'], sd['Final Predicted Fuel Consumption XGB After MC (kg/h)'],
                          label=name)
            self.ax2.plot(sd['Ship Speed (knots)'], sd['Final Predicted CO2 Emission XGB After MC (kg)'], label=name)
        self.ax1.plot(speeds, fc_all, color='black', linewidth=2, label='Combined')
        self.ax2.plot(speeds, co_all, color='black', linewidth=2, label='Combined')
        v = vals['Ship Speed (knots)']
        self.ax1.scatter(v, fcp, color='red', zorder=5, label='User')
        self.ax2.scatter(v, co2p, color='red', zorder=5, label='User')
        self.ax1.legend(fontsize=6);
        self.ax2.legend(fontsize=6)
        self.canvas.draw()

    def save_results(self):
        if not self.last:
            messagebox.showerror('Error', 'No prediction to save')
            return
        vals, fcp, co2p = self.last
        user_df = pd.DataFrame([{
            'Ship Speed (knots)': vals['Ship Speed (knots)'],
            'Final Predicted Fuel Consumption XGB After MC (kg/h)': fcp,
            'Final Predicted CO2 Emission XGB After MC (kg)': co2p
        }])
        try:
            with pd.ExcelWriter(output_path, engine='openpyxl', mode='a', if_sheet_exists='replace') as writer:
                user_df.to_excel(writer, sheet_name='User_Predictions', index=False)
            messagebox.showinfo('Saved', f'Results saved to {output_path}')
            self.b2.config(state='normal')
            self.b_save.config(state='disabled')
        except Exception as e:
            messagebox.showerror('Error', f'Failed to save: {e}')

    def generate_report(self):
        if not self.last: return
        vals, fcp, co2p = self.last
        # compute distances and scenario metrics
        scenario_metrics = {}
        for s in sheet_names:
            sd = excel_data[s]
            avg_fc = sd['Final Predicted Fuel Consumption XGB After MC (kg/h)'].mean()
            avg_co2 = sd['Final Predicted CO2 Emission XGB After MC (kg)'].mean()
            dist = np.hypot(fcp - avg_fc, co2p - avg_co2)
            scenario_metrics[s] = {'distance': dist, 'fc': avg_fc, 'co2': avg_co2}
        sorted_scenarios = sorted(scenario_metrics.items(), key=lambda x: x[1]['distance'])
        closest, second = sorted_scenarios[0], (sorted_scenarios[1] if len(sorted_scenarios) > 1 else (None, None))
        # suggestion
        min_fc_s = min(scenario_metrics.items(), key=lambda x: x[1]['fc'])[0]
        text = f"Input: {vals}\nPred FC: {fcp:.2f}, CO2: {co2p:.2f}\n"
        text += f"Closest: {closest[0]} (FC={closest[1]['fc']:.2f}, CO2={closest[1]['co2']:.2f})\n"
        if second[0]:
            text += f"Second Closest: {second[0]} (FC={second[1]['fc']:.2f}, CO2={second[1]['co2']:.2f})\n"
        text += f"Suggestion: improvements similar to {min_fc_s}."
        # update report box
        self.report_text.config(state='normal')
        self.report_text.delete('1.0', 'end')
        self.report_text.insert('end', text)
        self.report_text.config(state='disabled')
        # save report
        doc = Document()
        doc.add_heading('Prediction Report', 0)
        doc.add_paragraph(text)
        doc.save(report_path)
        messagebox.showinfo('Saved', f'Report saved to {report_path}')

    def zoom(self, factor):
        self.scale *= factor
        self.fig.set_size_inches(6 * self.scale, 6 * self.scale)
        self.canvas.draw()


if __name__ == '__main__':
    root = tk.Tk()
    PredictionApp(root)
    root.mainloop()
